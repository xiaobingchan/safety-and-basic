#pip install requests
#pip install beautifulsoup4
#pip install lxml
#pip install python-docx
#登录淘股吧
import requests
from bs4 import BeautifulSoup
from urllib.request import urlopen
import docx
import os
from docx import Document
data = {'userName': '18565453898', 'password': 'a12345678','checkCode':'','save':'Y','url':''}
r = requests.post('https://sso.taoguba.com.cn/web/login/submit', data=data)
print(r.text)
print(r.cookies.get_dict())
#登录老子博客第一页
page=15
#抓取每一页的所有链接地址。。9-12
list=[]
index=0
for i in range(13,page):
  r2 = requests.get('https://www.taoguba.com.cn/moreTopic?pageNum=15&pageNo='+str(i)+'&sortFlag=T&userID=53928', cookies=r.cookies)
  html = r2.text
  soup = BeautifulSoup(html, features='lxml')
  month = soup.find_all('a')
  for m in month: 
    url=m.get('href')
    url="https://www.taoguba.com.cn/"+url
    if url.find('Article')>=0:
      text = m.get_text().replace(u'\xa0', u' ')
      #print(text+","+url)
      #进去具体文章
      r3 = requests.get(url, cookies=r.cookies)
      html3 = r3.text
      soup2 = BeautifulSoup(html3, features='lxml')
      month2 = soup2.find('div', {"class": 'p_coten'})
      text2=month2.text
      text2 = text2.replace(u'\xa0', u' ')#p_tatime
      text2 = text2.replace(u' ', u'')
      time=soup2.find('span', {"class": 'p_tatime'})
      time=time.text
      time = time.replace(u'\xa0', u' ')

      # comment=[]
      # comment_all=soup2.find('div', {"class": 'lightenreply'})
      # for m in comment_all:
      #   #第一层
      #   soup_1=beautifulsoup(m, 'lxml')
      #   comment_1=soup_1.find('div', {"class": 'lightenreply'})
      #   soup3 = beautifulsoup(m, 'lxml')
      #   spans = soup3.find_all('a')
      #   if spans.text.find('老子')>=0:
      #       soup4 = beautifulsoup(m, 'lxml')
      #       p_div = soup4.find('div', {"class": 'pcnr_wz'})
      #       comment.append(p_div)

      artice={}
      artice['title']=text
      artice['time']=time
      artice['url']=url
      artice['content']=text2
      #artice['comment']=comment
      list.append(artice)
      index=index+1
      print(index)
#写入word文档
newfile="老子博客.docx"
if os.path.exists(newfile):
  os.remove(newfile)
document = Document()
for i in range(len(list)-1,-1,-1):
  document.add_heading(list[i]['title'], 2)
  document.add_paragraph(list[i]['time'])
  document.add_paragraph(list[i]['content'])
  document.add_paragraph(list[i]['url'])
  document.add_paragraph("")
document.save("老子博客.docx")  # 保存文档
#处理准备插入的图片问题
#爬评论


